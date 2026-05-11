"""
Obsidian Vault Reorganizer  (Taxonomy Edition v2)
=================================================
Classifies .md, .pdf, and .docx files in your vault and links them by topic.

Layers
------
1. PRIMARY TOPIC  (one per file) -> tag in frontmatter + section in MOC.md
2. CYBER DOMAINS  (multi-label)  -> copies into "cyber domains/<Domain>/"

PDFs / Word docs
----------------
We can't inject wikilinks inside binary files, so for each .pdf / .docx we
create a sidecar "<stem>.<ext>.index.md" that:
  - Lists tags + cyber domains in frontmatter
  - Embeds the original file with ![[file.pdf]]
  - Has the same auto-generated Related section as a normal note

Editable taxonomy
-----------------
The first run writes "taxonomy.json" next to this script. Edit that file to:
  - Rename a tag/topic (change the dict key)
  - Add or remove cyber domains
  - Tweak keywords (add new ones, drop noisy ones)
Re-run the script and the new taxonomy is used.

Commands
--------
    python obsidian_reorganize.py "D:\\bayyari\\bmt"                  # dry-run
    python obsidian_reorganize.py "D:\\bayyari\\bmt" --apply          # apply
    python obsidian_reorganize.py "D:\\bayyari\\bmt" --clean          # preview cleanup
    python obsidian_reorganize.py "D:\\bayyari\\bmt" --clean --apply  # remove all auto content
    python obsidian_reorganize.py --list-taxonomy                     # show current taxonomy
    python obsidian_reorganize.py --reset-taxonomy                    # rewrite taxonomy.json from defaults

Dependencies
------------
    pip install scikit-learn numpy pypdf python-docx
"""

import argparse
import json
import re
import shutil
import sys
from collections import defaultdict
from datetime import datetime
from pathlib import Path

try:
    from sklearn.feature_extraction.text import TfidfVectorizer
    from sklearn.metrics.pairwise import cosine_similarity
except ImportError:
    print("Missing required deps. Install with:\n    pip install scikit-learn numpy pypdf python-docx")
    sys.exit(1)


# =============================================================================
# DEFAULT TAXONOMY  -- written to taxonomy.json on first run, then editable
# =============================================================================

DEFAULT_PRIMARY_TOPICS = {
    'Linux'            : ['linux','ubuntu','debian','kali','centos','redhat','rhel','fedora','arch linux',
                          'bash','shell script','systemd','apt-get',' apt ',' yum',' dnf','grep ',' sed ',
                          ' awk ','/etc/','/var/',' sudo','cron','iptables','nftables'],
    'Windows'          : ['windows','win10','win11','win 10','win 11','powershell','wmi','sysmon',
                          'event viewer','active directory',' ad ','active-directory','ntlm','kerberos',
                          'registry','group policy',' gpo ','cmd.exe','windefend'],
    'Network'          : ['network','networking',' tcp','udp ','tcp/ip',' ip ',' dns','dhcp','vlan',
                          'subnet','routing','router','switch','firewall','vpn ',' ospf',' bgp',' nat ',
                          'packet','wireshark','nmap','port scan','osi model',' ftp ',' ssh '],
    'Cloud'            : ['cloud',' aws','amazon web services',' azure',' gcp','google cloud',' ec2',
                          ' s3 ','lambda ','kubernetes',' k8s','terraform','cloudformation','iam role',
                          ' vpc',' eks',' aks',' gke'],
    'HTU Camp'         : ['htu','htu camp','hashemite university','bayyari camp'],
    'Docker'           : ['docker','dockerfile','docker compose','docker-compose','container image',
                          'containerization','docker hub','podman'],
    'Tools'            : ['toolkit','utility','cheatsheet','cheat sheet','reference card','tooling',
                          'tool collection'],
    'Notion'           : ['notion','notion.so'],
    'N8N'              : [' n8n ','workflow automation','no-code automation','low-code automation'],
    'Certificates'     : ['certificate','certification','certified',' ceh ',' oscp','security+',
                          'sec+ ','comptia','sans ',' giac','gcih',' gpen','gcfa','exam','study guide',
                          'exam prep','cysa+','pentest+'],
    'Red Team'         : ['red team','red-team','offensive','exploit','exploitation','penetration test',
                          'pentest','pen test','attack chain','malware','payload',' c2 ',
                          'command and control','metasploit','cobalt strike','beacon','lateral movement',
                          'privilege escalation','privesc','reverse shell','bind shell','initial access'],
    'Blue Team'        : ['blue team','blue-team','defensive','defender',' siem',' edr',' xdr',
                          'detection engineering','threat hunt','log analysis','sigma rule','yara'],
    'DFIR'             : [' dfir','forensic','forensics','incident response','memory analysis',
                          'disk image','volatility','autopsy','timeline analysis',' ir ','triage',
                          'chain of custody'],
    'SOC'              : [' soc ','soc analyst','security operations center','tier 1','tier 2',
                          'tier-1','tier-2','alert triage','playbook','runbook'],
    'ISO27001'         : ['iso 27001','iso27001','iso-27001',' isms','annex a',
                          'statement of applicability'],
    'AI LLM Security'  : [' ai ','artificial intelligence',' llm','large language model',' gpt',
                          'prompt injection','jailbreak',' mcp ','model context protocol',' rag ',
                          'prompt engineering','ai security','ai safety'],
    'Phantom Scan'     : ['phantom scan','phantom-scan','phantomscan'],
    'ISC2 Org'         : ['isc2','isc-2','(isc)','cissp','ccsp',' sscp',' cgrc'],
    'Projects'         : ['my project','side project','build log','project plan','project notes'],
}

DEFAULT_CYBER_DOMAINS = {
    'Security Audit'                              : ['audit','auditing','audit trail','compliance audit',
                                                     'evidence collection','audit log','internal audit',
                                                     'external audit'],
    'Security and Risk Management'                : ['risk','governance',' grc','policy','compliance',
                                                     'gdpr','hipaa',' sox ','regulatory',
                                                     'business continuity',' bcp ',' drp',
                                                     'disaster recovery','risk assessment',
                                                     'risk management','threat model','risk register'],
    'Asset Security'                              : ['asset','asset management','data classification',
                                                     'classification','retention','data owner',
                                                     'data custodian','labeling','data lifecycle'],
    'Security Architecture and Engineering'       : ['architecture','secure design','secure architecture',
                                                     'cryptography','encryption','hashing',' pki ',
                                                     ' hsm',' tpm','zero trust','defense in depth',
                                                     'reference architecture'],
    'Communication and Network Security'          : ['network','firewall','vpn',' ids',' ips',' tls',
                                                     ' ssl','segmentation','vlan','packet','wireshark',
                                                     ' dns','dnssec','tcp/ip',' osi'],
    'Identity and Access Management (IAM)'        : [' iam ','identity','authentication','authorization',
                                                     ' sso',' mfa','2fa',' rbac',' abac','kerberos',
                                                     'ldap','oauth','saml','password policy',
                                                     'privileged access',' pam '],
    'Security Assessment and Testing'             : ['assessment','vulnerability scan',
                                                     'vulnerability assessment','pentest',
                                                     'penetration test','pen test','red team',
                                                     'code review','security testing','exploit'],
    'Security Operations'                         : [' soc ','security operations',' siem','incident',
                                                     'alert','log analysis',' edr',' xdr','detection',
                                                     'response',' dfir','forensic','incident response'],
    'Software Development Security'               : ['sdlc','devsecops','secure coding','owasp',' sast',
                                                     ' dast',' iast','code review','ci/cd',
                                                     'supply chain','dependency',' sbom'],
    'Cloud Security'                              : ['cloud',' aws',' azure',' gcp','kubernetes',
                                                     'container security','cspm','cwpp',' saas',
                                                     ' paas',' iaas','cloud iam'],
    'Endpoint Security'                           : ['endpoint',' edr','antivirus',' av ',
                                                     'host-based','workstation','laptop',
                                                     'mobile device',' mdm','hardening',
                                                     'endpoint protection'],
    'Threat Intelligence'                         : ['threat intel','threat intelligence',' ioc ',
                                                     'indicator of compromise',' ttp','mitre',
                                                     'att&ck','attack framework',' apt',
                                                     'threat hunting',' cti '],
    'Physical Security'                           : ['physical security','badge','surveillance',
                                                     'cctv','perimeter','door lock','biometric',
                                                     'facility','tailgating','mantrap'],
    'Data Loss Prevention (DLP)'                  : [' dlp ','data loss prevention',
                                                     'data exfiltration','exfiltration',
                                                     'data leakage','data leak','sensitive data',
                                                     'data masking'],
}

DEFAULT_FOLDER_RULES = {
    'exclude': [],
    'rename': {},
    'depth': 'top',
}

TAXONOMY_FILE = Path(__file__).parent / 'taxonomy.json'


def load_taxonomy(reset: bool = False):
    if reset or not TAXONOMY_FILE.exists():
        data = {
            'primary_topics': DEFAULT_PRIMARY_TOPICS,
            'cyber_domains': DEFAULT_CYBER_DOMAINS,
            'folder_rules': DEFAULT_FOLDER_RULES,
        }
        TAXONOMY_FILE.write_text(json.dumps(data, indent=2, ensure_ascii=False),
                                 encoding='utf-8')
        print(f'Wrote default taxonomy to: {TAXONOMY_FILE}')
        print('Edit that file to rename topics, add domains, or tweak keywords.\n')
    data = json.loads(TAXONOMY_FILE.read_text(encoding='utf-8'))
    folder_rules = data.get('folder_rules', {}) or {}
    raw_depth = folder_rules.get('depth', 'top')
    folder_rules = {
        'exclude': folder_rules.get('exclude', []) or [],
        'rename':  folder_rules.get('rename', {}) or {},
        'depth':   raw_depth if raw_depth in ('top', 'all') else 'top',
    }
    return data['primary_topics'], data['cyber_domains'], folder_rules


def save_taxonomy(primary_topics, cyber_domains, folder_rules):
    data = {
        'primary_topics': primary_topics,
        'cyber_domains': cyber_domains,
        'folder_rules': folder_rules,
    }
    TAXONOMY_FILE.write_text(json.dumps(data, indent=2, ensure_ascii=False),
                             encoding='utf-8')


def find_topic_or_domain(name, primary_topics, cyber_domains):
    """Case-insensitive lookup. Returns (kind, actual_name) or (None, None).
    kind is 'primary' or 'domain'. Primary topics win on conflict."""
    for k in primary_topics:
        if k.lower() == name.lower():
            return 'primary', k
    for k in cyber_domains:
        if k.lower() == name.lower():
            return 'domain', k
    return None, None


def edit_keywords(primary_topics, cyber_domains,
                  to_add, to_remove, to_show) -> bool:
    """Apply --add-keyword / --remove-keyword / --show-keywords from the CLI.
    Returns True if the taxonomy was modified."""
    changed = False

    for topic, keyword in (to_add or []):
        kind, actual = find_topic_or_domain(topic, primary_topics, cyber_domains)
        if not actual:
            print(f'! unknown topic or domain: "{topic}"')
            continue
        target = primary_topics if kind == 'primary' else cyber_domains
        if keyword in target[actual]:
            print(f'  "{keyword}" already in [{actual}]')
        else:
            target[actual].append(keyword)
            changed = True
            print(f'+ added "{keyword}" to [{actual}]')

    for topic, keyword in (to_remove or []):
        kind, actual = find_topic_or_domain(topic, primary_topics, cyber_domains)
        if not actual:
            print(f'! unknown topic or domain: "{topic}"')
            continue
        target = primary_topics if kind == 'primary' else cyber_domains
        if keyword not in target[actual]:
            print(f'  "{keyword}" not in [{actual}] (nothing to remove)')
        else:
            target[actual].remove(keyword)
            changed = True
            print(f'- removed "{keyword}" from [{actual}]')

    for topic in (to_show or []):
        kind, actual = find_topic_or_domain(topic, primary_topics, cyber_domains)
        if not actual:
            print(f'! unknown topic or domain: "{topic}"')
            continue
        target = primary_topics if kind == 'primary' else cyber_domains
        label  = 'primary topic' if kind == 'primary' else 'cyber domain'
        print(f'\n[{actual}]  ({label}, {len(target[actual])} keywords)')
        for kw in target[actual]:
            print(f'   {kw!r}')

    return changed


# =============================================================================
# regex / IO helpers
# =============================================================================

FRONTMATTER_RE = re.compile(r'^---\s*\n(.*?)\n---\s*\n', re.DOTALL)
AUTO_BEGIN     = '<!-- BMT_AUTO_BEGIN -->'
AUTO_END       = '<!-- BMT_AUTO_END -->'
# matches the auto block AND the optional `---` separator we put before it
AUTO_BLOCK_RE  = re.compile(
    r'\n*(?:---\s*\n+)?' + re.escape(AUTO_BEGIN) + r'.*?' + re.escape(AUTO_END) + r'\s*',
    re.DOTALL,
)
INVALID_FS_CHARS = re.compile(r'[<>:"|?*]')


def read_md(path: Path):
    text = path.read_text(encoding='utf-8', errors='replace')
    if text.startswith('\ufeff'):
        text = text[1:]
    m = FRONTMATTER_RE.match(text)
    if m:
        return m.group(1), text[m.end():]
    return '', text


def write_md(path: Path, frontmatter: str, body: str):
    if frontmatter.strip():
        out = f'---\n{frontmatter.strip()}\n---\n\n{body.lstrip()}'
    else:
        out = body
    path.write_text(out, encoding='utf-8')


def extract_pdf_text(path: Path, max_pages: int = 80) -> str:
    try:
        from pypdf import PdfReader
    except ImportError:
        try:
            from PyPDF2 import PdfReader  # legacy
        except ImportError:
            return ''
    try:
        reader = PdfReader(str(path))
        pages  = reader.pages[:max_pages]
        return '\n'.join((p.extract_text() or '') for p in pages)
    except Exception as e:
        print(f'  ! could not read PDF {path.name}: {e}')
        return ''


def extract_docx_text(path: Path) -> str:
    try:
        import docx  # python-docx
    except ImportError:
        return ''
    try:
        d = docx.Document(str(path))
        paras  = [p.text for p in d.paragraphs]
        tables = [' '.join(c.text for r in t.rows for c in r.cells) for t in d.tables]
        return '\n'.join(paras + tables)
    except Exception as e:
        print(f'  ! could not read DOCX {path.name}: {e}')
        return ''


def slugify(text: str) -> str:
    text = re.sub(r'[^\w\s-]', '', text, flags=re.UNICODE)
    text = re.sub(r'[\s_-]+', '-', text).strip('-')
    return text.lower() or 'untitled'


def safe_folder_name(name: str) -> str:
    return INVALID_FS_CHARS.sub('', name).strip().rstrip('.')


def sidecar_path(path: Path) -> Path:
    """Sidecar index for a non-md file: 'report.pdf' -> 'report.pdf.index.md'."""
    return path.parent / f'{path.name}.index.md'


def is_sidecar(path: Path) -> bool:
    return path.name.endswith('.index.md')


# =============================================================================
# classification
# =============================================================================

def score_keywords(text_lower: str, keywords) -> int:
    return sum(text_lower.count(kw.lower()) for kw in keywords)


def assign_primary(text_lower: str, primary_topics):
    scores = {topic: score_keywords(text_lower, kws)
              for topic, kws in primary_topics.items()}
    best = max(scores, key=scores.get)
    if scores[best] == 0:
        return 'Uncategorized'
    return best


def assign_domains(text_lower: str, cyber_domains, single_domain: bool = False):
    best = None
    best_score = 0
    matches = []
    for d, kws in cyber_domains.items():
        score = score_keywords(text_lower, kws)
        if score > 0:
            matches.append(d)
            if score > best_score:
                best = d
                best_score = score

    if not single_domain:
        return matches
    if best_score == 0:
        return []
    return [best]


# =============================================================================
# frontmatter helpers
# =============================================================================

def upsert_list_in_frontmatter(fm: str, key: str, values) -> str:
    values = [v for v in values if v]
    if not values:
        return fm
    if not fm.strip():
        return f'{key}: [{", ".join(values)}]'

    lines = fm.splitlines()
    for i, line in enumerate(lines):
        if line.strip().startswith(f'{key}:'):
            m_inline = re.match(rf'(\s*{key}:\s*)\[(.*?)\]\s*$', line)
            if m_inline:
                existing = [t.strip() for t in m_inline.group(2).split(',') if t.strip()]
                merged = existing + [v for v in values if v not in existing]
                lines[i] = f'{m_inline.group(1)}[{", ".join(merged)}]'
                return '\n'.join(lines)
            j = i + 1
            block, end = [], i
            while j < len(lines) and re.match(r'\s*-\s+', lines[j]):
                block.append(lines[j].strip().lstrip('-').strip())
                end = j
                j += 1
            merged = block + [v for v in values if v not in block]
            new_block = [f'{key}:'] + [f'  - {v}' for v in merged]
            lines[i:end + 1 if block else i + 1] = new_block
            return '\n'.join(lines)

    lines.append(f'{key}: [{", ".join(values)}]')
    return '\n'.join(lines)


def add_domain_tag_to_path(path: Path, domain: str, make_sidecars: bool = True) -> bool:
    """Add a domain slug to a file (md or pdf/docx sidecar). Returns True if changed."""
    domain_slug = slugify(domain)
    if path.suffix.lower() == '.md':
        fm, body = read_md(path)
        new_fm = upsert_list_in_frontmatter(fm, 'domains', [domain_slug])
        if new_fm != fm:
            write_md(path, new_fm, body)
            return True
        return False

    if path.suffix.lower() in ('.pdf', '.docx'):
        if not make_sidecars:
            return False
        side = sidecar_path(path)
        if side.exists():
            fm, body = read_md(side)
        else:
            fm, body = '', ''
        new_fm = upsert_list_in_frontmatter(fm, 'domains', [domain_slug])
        new_fm = upsert_list_in_frontmatter(new_fm, 'source', [path.name])
        if side.exists() and new_fm == fm:
            return False
        if not body.strip():
            body = (
                f'# {path.stem}\n\n'
                f'_Index for `{path.name}`._\n\n'
                f'![[{path.name}]]\n'
            )
        write_md(side, new_fm, body)
        return True

    return False


def remove_known_from_tags(fm: str, known_slugs: set) -> str:
    """Drop only our slugs from `tags:`. Remove the line if it goes empty."""
    if not fm.strip():
        return fm
    lines = fm.splitlines()
    out = []
    i = 0
    while i < len(lines):
        line = lines[i]
        if line.strip().startswith('tags:'):
            m = re.match(r'(\s*tags:\s*)\[(.*?)\]\s*$', line)
            if m:
                vals = [t.strip() for t in m.group(2).split(',') if t.strip()]
                kept = [v for v in vals if v not in known_slugs]
                if kept:
                    out.append(f'{m.group(1)}[{", ".join(kept)}]')
                # else drop the line
                i += 1
                continue
            # block form
            block_idx = []
            j = i + 1
            while j < len(lines) and re.match(r'\s*-\s+', lines[j]):
                block_idx.append(j)
                j += 1
            if block_idx:
                kept_lines = []
                for k in block_idx:
                    val = lines[k].strip().lstrip('-').strip()
                    if val not in known_slugs:
                        kept_lines.append(lines[k])
                if kept_lines:
                    out.append(line)
                    out.extend(kept_lines)
                i = j
                continue
        out.append(line)
        i += 1
    return '\n'.join(out)


def remove_line(fm: str, key: str) -> str:
    """Remove `key:` and any block-style children below it."""
    if not fm.strip():
        return fm
    lines = fm.splitlines()
    out = []
    i = 0
    while i < len(lines):
        line = lines[i]
        if line.strip().startswith(f'{key}:'):
            i += 1
            while i < len(lines) and re.match(r'\s*-\s+', lines[i]):
                i += 1
            continue
        out.append(line)
        i += 1
    return '\n'.join(out)


# =============================================================================
# discovery & classification of all files
# =============================================================================

def discover(vault: Path, cyber_root: Path):
    """Yield md/pdf/docx files we should classify (excluding our own outputs)."""
    for p in vault.rglob('*'):
        if not p.is_file():
            continue
        if cyber_root in p.parents:
            continue
        if p.suffix.lower() not in ('.md', '.pdf', '.docx'):
            continue
        if p.name == 'MOC.md':
            continue
        if is_sidecar(p):
            continue
        if p.name.startswith('_'):
            continue
        yield p


def build_haystack(path: Path) -> str:
    if path.suffix.lower() == '.md':
        _, body = read_md(path)
        body = AUTO_BLOCK_RE.sub('', body)
        return f' {path.stem} {body} '.lower()
    if path.suffix.lower() == '.pdf':
        return f' {path.stem} {extract_pdf_text(path)} '.lower()
    if path.suffix.lower() == '.docx':
        return f' {path.stem} {extract_docx_text(path)} '.lower()
    return f' {path.stem} '.lower()


def _norm(s: str) -> str:
    return (s or '').strip().lower()


def folder_parts(path: Path, vault: Path, depth: str) -> list:
    rel = path.relative_to(vault)
    parts = list(rel.parent.parts)
    return parts[:1]


def folder_tag_slugs(path: Path, vault: Path, folder_rules) -> list:
    """Return unique, slugified parent folder names for a file."""
    depth = folder_rules.get('depth', 'all')
    exclude = {_norm(x) for x in folder_rules.get('exclude', [])}
    rename = { _norm(k): v for k, v in (folder_rules.get('rename', {}) or {}).items() }

    out = []
    for part in folder_parts(path, vault, depth):
        if not part or part.startswith('.') or part.startswith('_'):
            continue
        if _norm(part) in exclude:
            continue
        tag_source = rename.get(_norm(part), part)
        slug = slugify(tag_source)
        if slug == 'untitled':
            continue
        if slug not in out:
            out.append(slug)
    return out


def list_folder_counts(files, vault: Path, folder_rules) -> dict:
    depth = folder_rules.get('depth', 'all')
    exclude = {_norm(x) for x in folder_rules.get('exclude', [])}
    counts = {}
    for p in files:
        for part in folder_parts(p, vault, depth):
            if not part or part.startswith('.') or part.startswith('_'):
                continue
            if _norm(part) in exclude:
                continue
            name = part.strip()
            if not name:
                continue
            counts[name] = counts.get(name, 0) + 1
    return dict(sorted(counts.items(), key=lambda x: (-x[1], x[0].lower())))


def suggest_folder_domains(files, vault: Path, cyber_domains, folder_rules) -> dict:
    existing = {k.lower() for k in cyber_domains}
    counts = list_folder_counts(files, vault, folder_rules)
    return {k: v for k, v in counts.items() if k.lower() not in existing}


# =============================================================================
# CLEAN mode
# =============================================================================

def _check_cancel(cancel_event) -> bool:
    if cancel_event and cancel_event.is_set():
        print('\nCanceled by user.')
        return True
    return False


def do_clean(vault: Path, cyber_root: Path, primary_topics, cyber_domains,
             folder_rules, apply: bool, remove_folder_tags: bool,
             cancel_event=None):
    known_topic_slugs = {slugify(k) for k in primary_topics}
    known_domain_slugs = {slugify(k) for k in cyber_domains}
    folder_slugs = set()
    if remove_folder_tags:
        for p in discover(vault, cyber_root):
            folder_slugs.update(folder_tag_slugs(p, vault, folder_rules))
    all_known = known_topic_slugs | known_domain_slugs | folder_slugs

    md_to_strip = []
    sidecars_to_delete = []
    for p in vault.rglob('*'):
        if _check_cancel(cancel_event):
            return
        if not p.is_file() or cyber_root in p.parents:
            continue
        if p.suffix.lower() != '.md':
            continue
        text = p.read_text(encoding='utf-8', errors='replace')
        if is_sidecar(p) and AUTO_BEGIN in text:
            sidecars_to_delete.append(p)
        elif AUTO_BEGIN in text or 'tags:' in text or 'domains:' in text:
            md_to_strip.append(p)

    moc = vault / 'MOC.md'

    print('=' * 72)
    print('CLEAN PLAN')
    print('=' * 72)
    print(f'\nWill strip auto-block, our tag slugs, and our `domains:` line from')
    print(f'{len(md_to_strip)} .md file(s).')
    if remove_folder_tags:
        print('Also removing folder-based tag slugs from `tags:`.')
    print(f'\nWill delete {len(sidecars_to_delete)} sidecar(s) (*.index.md):')
    for p in sidecars_to_delete[:10]:
        print(f'   - {p.relative_to(vault)}')
    if len(sidecars_to_delete) > 10:
        print(f'   ... and {len(sidecars_to_delete) - 10} more')

    print(f'\nWill delete: {moc}' if moc.exists() else f'\n(MOC.md not present)')
    print(f'Will delete folder: {cyber_root}' if cyber_root.exists()
          else f'(cyber-domains folder not present)')

    if not apply:
        print('\nDry-run only. Re-run with --clean --apply to actually remove.')
        return

    if _check_cancel(cancel_event):
        return

    # backup
    backup = vault.parent / f'{vault.name}_backup_{datetime.now():%Y%m%d_%H%M%S}'
    shutil.copytree(vault, backup)
    print(f'\nBackup created: {backup}')

    for p in md_to_strip:
        if _check_cancel(cancel_event):
            return
        fm, body = read_md(p)
        body = AUTO_BLOCK_RE.sub('', body).rstrip() + '\n'
        fm = remove_known_from_tags(fm, all_known)
        fm = remove_line(fm, 'domains')
        write_md(p, fm, body)
    for p in sidecars_to_delete:
        if _check_cancel(cancel_event):
            return
        p.unlink()
    if moc.exists():
        moc.unlink()
    if cyber_root.exists():
        shutil.rmtree(cyber_root)
    print('Cleanup done.')


def remove_folder_tags_from_notes(vault: Path, cyber_root: Path, slugs: list, apply: bool):
    if not slugs:
        print('No folder tag slugs provided.')
        return
    md_files = [p for p in vault.rglob('*.md') if cyber_root not in p.parents]
    targets = []
    for p in md_files:
        fm, body = read_md(p)
        new_fm = remove_known_from_tags(fm, set(slugs))
        if new_fm != fm:
            targets.append((p, new_fm, body))

    print('=' * 72)
    print('REMOVE FOLDER TAGS')
    print('=' * 72)
    print(f'Will remove {len(slugs)} tag(s) from {len(targets)} file(s).')
    for p, _, _ in targets[:10]:
        print(f'   - {p.relative_to(vault)}')
    if len(targets) > 10:
        print(f'   ... and {len(targets) - 10} more')

    if not apply:
        print('\nDry-run only. Re-run with --apply to write changes.')
        return

    backup = vault.parent / f'{vault.name}_backup_{datetime.now():%Y%m%d_%H%M%S}'
    shutil.copytree(vault, backup)
    print(f'\nBackup created: {backup}')

    for p, new_fm, body in targets:
        write_md(p, new_fm, body)
    print('Folder tag removal done.')


def apply_domain_to_files(vault: Path, file_paths: list, domain: str,
                          make_sidecars: bool = True) -> dict:
    """Apply a single domain tag to a list of vault-relative file paths."""
    changed = 0
    skipped = 0
    for rel in file_paths:
        try:
            path = (vault / rel).resolve()
        except Exception:
            skipped += 1
            continue
        if not path.exists() or not path.is_file():
            skipped += 1
            continue
        if add_domain_tag_to_path(path, domain, make_sidecars):
            changed += 1
        else:
            skipped += 1
    return {'changed': changed, 'skipped': skipped}


# =============================================================================
# APPLY mode
# =============================================================================

def write_md_note(path: Path, primary, domains, related_lines, folder_tags,
                  primary_topics, cyber_domains):
    """Rewrite a .md note: refresh tags/domains in frontmatter and the auto block."""
    fm, body = read_md(path)
    body = AUTO_BLOCK_RE.sub('', body).rstrip()

    block = [AUTO_BEGIN, '', '## Related', '']
    block += related_lines or ['_No strongly related notes found._']
    block += ['', f'**Topic:** [[MOC#{primary}|{primary}]]']
    if domains:
        dom_links = ', '.join(f'[[Cyber Domains MOC#{d}|{d}]]' for d in domains)
        block.append(f'**Cyber Domains:** {dom_links}')
    block += ['', AUTO_END]

    new_body = body + '\n\n---\n\n' + '\n'.join(block) + '\n'
    tags = [slugify(primary)] + (folder_tags or [])
    fm = upsert_list_in_frontmatter(fm, 'tags', tags)
    fm = upsert_list_in_frontmatter(fm, 'domains', [slugify(d) for d in domains])
    write_md(path, fm, new_body)


def write_sidecar(side: Path, source: Path, primary, domains, related_lines,
                  folder_tags):
    """Sidecar index for a PDF/DOCX with embed + tags + Related."""
    tags = [slugify(primary)] + (folder_tags or [])
    fm = upsert_list_in_frontmatter('', 'tags', tags)
    fm = upsert_list_in_frontmatter(fm, 'domains', [slugify(d) for d in domains])
    fm = upsert_list_in_frontmatter(fm, 'source', [source.name])

    block = [AUTO_BEGIN, '', '## Related', '']
    block += related_lines or ['_No strongly related notes found._']
    block += ['', f'**Topic:** [[MOC#{primary}|{primary}]]']
    if domains:
        dom_links = ', '.join(f'[[Cyber Domains MOC#{d}|{d}]]' for d in domains)
        block.append(f'**Cyber Domains:** {dom_links}')
    block += ['', AUTO_END]

    body = (
        f'# {source.stem}\n\n'
        f'_Index for `{source.name}`._\n\n'
        f'![[{source.name}]]\n\n'
        f'---\n\n' + '\n'.join(block) + '\n'
    )
    write_md(side, fm, body)


def link_for(note) -> str:
    """How to link to a note from another note's Related section."""
    p = note['path']
    if p.suffix.lower() == '.md':
        return f'- [[{p.stem}]]'
    return f'- [[{p.name}]]   _(via [[{sidecar_path(p).stem}]])_'


def do_apply(vault: Path, cyber_root: Path, primary_topics, cyber_domains,
             folder_rules, top_k: int, threshold: float, make_sidecars: bool,
             apply: bool, cyber_folder_name: str, add_folder_tags: bool,
             suggest_domains: bool, single_domain: bool, cancel_event=None):

    files = list(discover(vault, cyber_root))
    if not files:
        print('No .md, .pdf, or .docx files found.')
        return

    print(f'Scanning {len(files)} file(s)...')
    notes = []
    for p in files:
        if _check_cancel(cancel_event):
            return
        haystack = build_haystack(p)
        primary  = assign_primary(haystack, primary_topics)
        domains  = assign_domains(haystack, cyber_domains, single_domain)
        folder_tags = folder_tag_slugs(p, vault, folder_rules) if add_folder_tags else []
        notes.append({'path': p, 'haystack': haystack,
                      'primary': primary, 'domains': domains,
                      'folder_tags': folder_tags})

    by_primary = defaultdict(list)
    by_domain  = defaultdict(list)
    for n in notes:
        by_primary[n['primary']].append(n)
        for d in n['domains']:
            by_domain[d].append(n)

    # TF-IDF for "Related"
    corpus = [f'{n["path"].stem} {n["path"].stem} {n["haystack"]}' for n in notes]
    vec = TfidfVectorizer(max_features=4000, ngram_range=(1, 2), min_df=1,
                          token_pattern=r"(?u)\b\w[\w'-]+\b")
    X   = vec.fit_transform(corpus)
    sim = cosine_similarity(X)
    related_idx = []
    for i in range(sim.shape[0]):
        if _check_cancel(cancel_event):
            return
        sim[i, i] = -1
        idx = sim[i].argsort()[::-1][:top_k]
        idx = [j for j in idx if sim[i, j] >= threshold]
        related_idx.append(idx)

    if suggest_domains:
        suggestions = suggest_folder_domains(files, vault, cyber_domains, folder_rules)
        print(f'\n{"=" * 72}\nFOLDER DOMAIN SUGGESTIONS\n{"=" * 72}')
        if suggestions:
            for name, count in suggestions.items():
                print(f'  - {name}  ({count} file(s))')
        else:
            print('  (no new folder names found)')

    # ---- print plan ---------------------------------------------------------
    print(f'\n{"=" * 72}\nPRIMARY TOPIC PLAN\n{"=" * 72}')
    for topic in list(primary_topics) + ['Uncategorized']:
        ns = by_primary.get(topic, [])
        if not ns:
            continue
        print(f'\n[{topic}]  ({len(ns)} files)')
        for n in sorted(ns, key=lambda x: x['path'].name.lower()):
            kind = n['path'].suffix.lower().lstrip('.')
            print(f'   - [{kind}] {n["path"].relative_to(vault)}')

    print(f'\n{"=" * 72}\nCYBER DOMAIN ASSIGNMENTS  (multi-label)\n{"=" * 72}')
    for domain in cyber_domains:
        ns = by_domain.get(domain, [])
        if not ns:
            continue
        print(f'\n[{domain}]  ({len(ns)} files)')
        for n in sorted(ns, key=lambda x: x['path'].name.lower()):
            extras = [d for d in n['domains'] if d != domain]
            tail = f'  (also: {", ".join(extras)})' if extras else ''
            print(f'   - {n["path"].name}{tail}')

    no_domain = [n for n in notes if not n['domains']]
    if no_domain:
        print(f'\n[no cyber-domain match]  ({len(no_domain)} files)')
        for n in sorted(no_domain, key=lambda x: x['path'].name.lower()):
            print(f'   - {n["path"].name}')

    if not apply:
        print('\nDry-run only. Re-run with --apply to write changes.')
        return

    if _check_cancel(cancel_event):
        return

    # ---- backup -------------------------------------------------------------
    backup = vault.parent / f'{vault.name}_backup_{datetime.now():%Y%m%d_%H%M%S}'
    shutil.copytree(vault, backup)
    print(f'\nBackup created: {backup}')

    # ---- write each note / sidecar ------------------------------------------
    for i, note in enumerate(notes):
        if _check_cancel(cancel_event):
            return
        related = [link_for(notes[j]) for j in related_idx[i]]
        if note['path'].suffix.lower() == '.md':
            write_md_note(note['path'], note['primary'], note['domains'], related,
                          note['folder_tags'],
                          primary_topics, cyber_domains)
        elif make_sidecars:
            write_sidecar(sidecar_path(note['path']), note['path'],
                          note['primary'], note['domains'], related,
                          note['folder_tags'])

    # ---- main MOC -----------------------------------------------------------
    moc = vault / 'MOC.md'
    lines = ['# Map of Content', '',
             f'_Auto-generated on {datetime.now():%Y-%m-%d %H:%M}._', '',
             f'{len(notes)} files. See also '
             f'[[{cyber_folder_name}/Cyber Domains MOC|Cyber Domains MOC]].', '']
    for topic in list(primary_topics) + ['Uncategorized']:
        ns = by_primary.get(topic, [])
        if not ns:
            continue
        lines += [f'## {topic}', '']
        for n in sorted(ns, key=lambda x: x['path'].name.lower()):
            p = n['path']
            link = f'[[{p.stem}]]' if p.suffix.lower() == '.md' else f'[[{p.name}]]'
            lines.append(f'- {link}')
        lines.append('')
    moc.write_text('\n'.join(lines), encoding='utf-8')
    print(f'Wrote: {moc}')

    # ---- cyber domains folder with COPIES -----------------------------------
    if cyber_root.exists():
        for sub in cyber_root.iterdir():
            if sub.is_dir():
                shutil.rmtree(sub)
            elif sub.name == 'Cyber Domains MOC.md':
                sub.unlink()
    cyber_root.mkdir(exist_ok=True)

    cyber_lines = ['# Cyber Domains MOC', '',
                   f'_Auto-generated on {datetime.now():%Y-%m-%d %H:%M}._', '',
                   'Files appear in every domain they match. '
                   'See also [[../MOC|Main MOC]].', '']
    for domain in cyber_domains:
        if _check_cancel(cancel_event):
            return
        ns = by_domain.get(domain, [])
        if not ns:
            continue
        folder = cyber_root / safe_folder_name(domain)
        folder.mkdir(exist_ok=True)
        for n in ns:
            if _check_cancel(cancel_event):
                return
            try:
                shutil.copy2(n['path'], folder / n['path'].name)
            except Exception as e:
                print(f'  ! could not copy {n["path"].name}: {e}')
        cyber_lines += [f'## {domain}', '',
                        f'_Folder: `{cyber_folder_name}/{safe_folder_name(domain)}/`_', '']
        for n in sorted(ns, key=lambda x: x['path'].name.lower()):
            p = n['path']
            link = f'[[{p.stem}]]' if p.suffix.lower() == '.md' else f'[[{p.name}]]'
            cyber_lines.append(f'- {link}')
        cyber_lines.append('')
    (cyber_root / 'Cyber Domains MOC.md').write_text(
        '\n'.join(cyber_lines), encoding='utf-8')
    print(f'Wrote: {cyber_root / "Cyber Domains MOC.md"}')

    print('\nDone. Open the vault in Obsidian.')
    print(f'(Edit "{TAXONOMY_FILE.name}" anytime to rename topics or tweak keywords.)')


# =============================================================================
# main
# =============================================================================

def main():
    ap = argparse.ArgumentParser()
    ap.add_argument('vault', nargs='?', help=r'Path to the vault, e.g. D:\bayyari\bmt')
    ap.add_argument('--apply', action='store_true', help='Actually write changes.')
    ap.add_argument('--clean', action='store_true',
                    help='Remove all auto-generated content (MOC, domain copies, '
                         'sidecars, auto-block, our tag slugs).')
    ap.add_argument('--no-sidecars', action='store_true',
                    help='Do not create sidecar .index.md files for PDFs/DOCX.')
    ap.add_argument('--top-k-related', type=int, default=5)
    ap.add_argument('--threshold', type=float, default=0.10)
    ap.add_argument('--cyber-folder', default='cyber domains',
                    help='Folder name (under the vault) for domain copies.')
    ap.add_argument('--folder-tags', action='store_true',
                    help='Add parent folder names as tag slugs.')
    ap.add_argument('--suggest-folder-domains', action='store_true',
                    help='Print folder names not present in cyber domains.')
    ap.add_argument('--single-domain', action='store_true',
                    help='Assign at most one cyber domain (highest score).')
    ap.add_argument('--folder-depth', choices=['all', 'top'],
                    help='Use all parent folders or only the top-level folder.')
    ap.add_argument('--list-folders', action='store_true',
                    help='List detected folders and counts, then exit.')
    ap.add_argument('--add-folder-domains', nargs='*',
                    help='Add folder names as cyber domains. If empty, add all detected.')
    ap.add_argument('--exclude-folder-tag', action='append', metavar='NAME',
                    help='Exclude a folder name from folder-based tags. Can be repeated.')
    ap.add_argument('--rename-folder-tag', nargs=2, action='append',
                    metavar=('OLD', 'NEW'),
                    help='Rename a folder-based tag. Can be repeated.')
    ap.add_argument('--clear-folder-tag-rules', action='store_true',
                    help='Clear folder tag rename/exclude rules.')
    ap.add_argument('--remove-folder-tags', action='store_true',
                    help='Remove all folder-based tag slugs from notes.')
    ap.add_argument('--remove-folder-tag', action='append', metavar='NAME',
                    help='Remove a specific folder tag slug from notes. Can be repeated.')
    ap.add_argument('--list-taxonomy', action='store_true',
                    help='Print current taxonomy.json and exit.')
    ap.add_argument('--reset-taxonomy', action='store_true',
                    help='Overwrite taxonomy.json with the built-in defaults and exit.')
    ap.add_argument('--add-keyword', nargs=2, action='append',
                    metavar=('TOPIC', 'KEYWORD'),
                    help='Add KEYWORD to a primary topic or cyber domain. '
                         'Can be repeated. Example: '
                         '--add-keyword "Linux" "ubuntu" --add-keyword "Cloud" "azure ad"')
    ap.add_argument('--remove-keyword', nargs=2, action='append',
                    metavar=('TOPIC', 'KEYWORD'),
                    help='Remove KEYWORD from a primary topic or cyber domain. '
                         'Can be repeated.')
    ap.add_argument('--show-keywords', action='append', metavar='TOPIC',
                    help='Print all keywords for a topic or domain. Can be repeated.')
    args = ap.parse_args()

    if args.reset_taxonomy:
        load_taxonomy(reset=True)
        return

    primary_topics, cyber_domains, folder_rules = load_taxonomy()

    if args.list_taxonomy:
        print(f'Taxonomy file: {TAXONOMY_FILE}\n')
        print('PRIMARY TOPICS  (one per file):')
        for k, v in primary_topics.items():
            print(f'  - {k}   [{len(v)} keywords]')
        print('\nCYBER DOMAINS  (multi-label, become folders):')
        for k, v in cyber_domains.items():
            print(f'  - {k}   [{len(v)} keywords]')
        print('\nFOLDER RULES:')
        print(f'  - depth: {folder_rules.get("depth", "all")}')
        print(f'  - excluded: {len(folder_rules.get("exclude", []))}')
        print(f'  - renamed: {len(folder_rules.get("rename", {}))}')
        return

    # --- keyword edits (modify taxonomy.json) --------------------------------
    folder_rules_changed = False

    if args.clear_folder_tag_rules:
        folder_rules['exclude'] = []
        folder_rules['rename'] = {}
        folder_rules_changed = True

    if args.exclude_folder_tag:
        for name in args.exclude_folder_tag:
            if name not in folder_rules['exclude']:
                folder_rules['exclude'].append(name)
                folder_rules_changed = True

    if args.rename_folder_tag:
        for old, new in args.rename_folder_tag:
            folder_rules['rename'][old] = new
        folder_rules_changed = True

    if args.folder_depth:
        folder_rules['depth'] = args.folder_depth
        folder_rules_changed = True

    if args.add_keyword or args.remove_keyword or args.show_keywords:
        changed = edit_keywords(primary_topics, cyber_domains,
                                args.add_keyword, args.remove_keyword,
                                args.show_keywords)
        if changed:
            save_taxonomy(primary_topics, cyber_domains, folder_rules)
            print(f'\nSaved: {TAXONOMY_FILE}')
        # If only keyword commands were requested (no vault), exit here.
        if not args.vault:
            return

    if folder_rules_changed:
        save_taxonomy(primary_topics, cyber_domains, folder_rules)
        print(f'\nSaved: {TAXONOMY_FILE}')
        if not args.vault:
            return

    if not args.vault:
        ap.error('vault path is required unless using '
                 '--list-taxonomy / --reset-taxonomy / --show-keywords / '
                 '--add-keyword / --remove-keyword')
    vault = Path(args.vault)
    if not vault.is_dir():
        print(f'ERROR: {vault} is not a directory.')
        sys.exit(1)
    cyber_root = vault / args.cyber_folder

    if args.list_folders:
        files = list(discover(vault, cyber_root))
        counts = list_folder_counts(files, vault, folder_rules)
        print(f'\nDetected folders (depth={folder_rules.get("depth", "all")})')
        if counts:
            for name, count in counts.items():
                print(f'  - {name}  ({count} file(s))')
        else:
            print('  (no folders found)')
        return

    if args.add_folder_domains is not None:
        files = list(discover(vault, cyber_root))
        counts = suggest_folder_domains(files, vault, cyber_domains, folder_rules)
        to_add = args.add_folder_domains or list(counts.keys())
        added = []
        for name in to_add:
            if name not in cyber_domains:
                cyber_domains[name] = []
                added.append(name)
        if added:
            save_taxonomy(primary_topics, cyber_domains, folder_rules)
            print(f'Added {len(added)} domain(s) from folders.')
            print(f'Saved: {TAXONOMY_FILE}')
        else:
            print('No new folder domains to add.')

    if args.remove_folder_tags or args.remove_folder_tag:
        to_remove = []
        if args.remove_folder_tags:
            files = list(discover(vault, cyber_root))
            slug_set = set()
            for p in files:
                slug_set.update(folder_tag_slugs(p, vault, folder_rules))
            to_remove = sorted(slug_set)
        if args.remove_folder_tag:
            to_remove += [slugify(x) for x in args.remove_folder_tag]
        to_remove = sorted({s for s in to_remove if s})
        remove_folder_tags_from_notes(vault, cyber_root, to_remove, args.apply)
        return

    if args.clean:
        do_clean(vault, cyber_root, primary_topics, cyber_domains,
                 folder_rules, args.apply, args.folder_tags)
    else:
        do_apply(vault, cyber_root, primary_topics, cyber_domains,
                 folder_rules, args.top_k_related, args.threshold,
                 not args.no_sidecars, args.apply, args.cyber_folder,
                 args.folder_tags, args.suggest_folder_domains,
                 args.single_domain)


if __name__ == '__main__':
    main()